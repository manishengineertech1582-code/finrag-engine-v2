"""
DOCX Loader
============
Extracts text from Word documents paragraph by paragraph.
Groups paragraphs into logical sections (separated by headings).

Metadata per section:
  source, section, section_index, paragraph_index, doc_type="docx", user_id
"""

import logging
import os
from typing import List

from langchain_core.documents import Document

logger = logging.getLogger(__name__)


def _get_style_name(para) -> str:
    """Safely extract paragraph style name; returns '' if style is missing."""
    try:
        if para.style is not None and hasattr(para.style, "name") and para.style.name:
            return para.style.name
    except Exception:
        pass
    return ""


def load_docx(file_path: str, user_id: str | None = None) -> List[Document]:
    """
    Load a DOCX file into LangChain Documents grouped by section.

    Args:
        file_path: Path to the .docx file.
        user_id:   Tenant / user identifier.

    Returns:
        List of Documents, one per logical section.
    """
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"DOCX not found: {file_path}")

    filename = os.path.basename(file_path)

    try:
        from docx import Document as DocxDocument

        doc = DocxDocument(file_path)
        logger.info("DOCX opened | paragraphs=%d | file=%s", len(doc.paragraphs), filename)

        sections: List[dict] = []
        current_heading = "Introduction"
        current_text: List[str] = []
        section_index = 0
        skipped = 0

        for para_index, para in enumerate(doc.paragraphs):
            try:
                text = para.text.strip() if para.text else ""
                if not text:
                    continue

                style_name = _get_style_name(para)

                # Detect headings to split sections
                if style_name.startswith("Heading"):
                    # Save previous section
                    if current_text:
                        sections.append({
                            "heading": current_heading,
                            "text": "\n".join(current_text),
                            "index": section_index,
                            "paragraph_index": para_index,
                        })
                        section_index += 1

                    current_heading = text
                    current_text = []
                else:
                    current_text.append(text)

            except Exception as e:
                skipped += 1
                logger.warning("Skipping paragraph %d due to error: %s", para_index, e)
                continue

        if skipped:
            logger.warning("DOCX skipped %d malformed paragraphs | file=%s", skipped, filename)

        # Save final section
        if current_text:
            sections.append({
                "heading": current_heading,
                "text": "\n".join(current_text),
                "index": section_index,
                "paragraph_index": -1,
            })

        documents = [
            Document(
                page_content=s["text"],
                metadata={
                    "source": filename,
                    "section": s["heading"],
                    "page_or_sheet": f"section_{s['index']}",
                    "paragraph_index": s["paragraph_index"],
                    "doc_type": "docx",
                    "user_id": user_id,
                }
            )
            for s in sections
            if s["text"].strip()
        ]

        logger.info("DOCX loaded | sections=%d | file=%s", len(documents), filename)
        return documents

    except (FileNotFoundError, RuntimeError):
        raise
    except Exception as e:
        logger.exception("Failed to load DOCX: %s", file_path)
        raise RuntimeError(f"DOCX load failed: {filename}") from e
